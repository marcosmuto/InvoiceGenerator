from docx import Document

from UpdateInvoiceContent import getNewInvoiceDescription, updateInvoiceNumber

# Wont create unit test for this function as it uses out of the box docx elements
# also the dependable functions to update the content as covered            
def updateInvoice(documentFileName):
    
    invoice = Document(documentFileName)

    for table in invoice.tables:
        for row in table.rows:
            for cell in row.cells:

                if cell.text.startswith("INVOICE"):
                    #The invoice cell have 3 paragraphs:
                    # 1. the INVOICE title
                    # 2. the invoice number like Invoice #123
                    # 3. the invoice date
                    newText = updateInvoiceNumber(cell.paragraphs[1].text)
                    cell.paragraphs[1].text = newText

                if cell.text.startswith("Contractor"):
                    # need to update the paragraph on the cell to preserve the formatting
                    newText = getNewInvoiceDescription(cell.paragraphs[0].text)
                    cell.paragraphs[0].text = cell.text = newText


    invoice.save(documentFileName)

