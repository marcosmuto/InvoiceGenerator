from docx import Document

from InvoiceContentHelper import InvoiceContentHelper

class InvoiceContentUpdate:

    contentHelper = InvoiceContentHelper()

    def __init__(self, invoice_file_name, invoice_values_file):
        self.__invoice_file_name = invoice_file_name
        self.__invoice_values_file = invoice_values_file

    # Wont create unit test for this function as it uses out of the box docx elements
    # also the dependable functions to update the content are covered            
    def UpdateInvoiceContent(self):
        
        invoice = Document(self.__invoice_file_name)

        invoice_values = self.contentHelper.LoadInvoiceValues(self.__invoice_values_file)
        # store the name of the person for the next invoice value
        current_invoice_name = ""

        for table in invoice.tables:
            for row in table.rows:
                for cell in row.cells:

                    if cell.text.startswith("INVOICE"):
                        #The invoice cell have 3 paragraphs:
                        # 1. the INVOICE title
                        # 2. the invoice number like Invoice #123
                        # 3. the invoice date
                        newText = self.contentHelper.UpdateInvoiceNumber(cell.paragraphs[1].text)
                        cell.paragraphs[1].text = newText
                        print("Updated invoice number: " + newText)

                    if cell.text.startswith("Contractor"):
                        # need to update the paragraph on the cell to preserve the formatting
                        newText = self.contentHelper.GetNewInvoiceDescription(cell.paragraphs[0].text)
                        cell.paragraphs[0].text = newText
                        # there is no direct link from the name and the invoice cell value
                        current_invoice_name = self.contentHelper.GetContractorName(newText)

                    if cell.text.startswith("SUBTOTAL"):
                        current_invoice_name = "Total"

                    if cell.text.startswith("$"):
                        invoice_value = "${:,.2f}".format(invoice_values[current_invoice_name].value)
                        cell.paragraphs[0].text = invoice_value
                        print("Updated invoice value {:>25} - {}".format(current_invoice_name, invoice_value))

        invoice.save(self.__invoice_file_name)

